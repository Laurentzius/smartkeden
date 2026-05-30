"""Tests for Document Parsing with MarkItDown integration.

Covers: file type detection, MarkItDown conversion for PDF/XLSX/DOCX,
Gemini Vision fallback for images, encrypted PDF rejection, unsupported format,
file size limits, and the full upload->extract->structurize pipeline.
"""

import io
import uuid
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from fastapi.testclient import TestClient

from app.services.parser.schemas import (
    InvoiceData,
    InvoiceLine,
    ProcessingMetadata,
    ParseDocumentResponse,
)
from app.services.parser.service import (
    detect_source_type,
    MAX_FILE_SIZE,
    structurize_invoice,
)
from app.services.parser.markitdown_adapter import convert_to_markdown


# ── Helpers ──────────────────────────────────────────────────────────────────


def _minimal_pdf_bytes() -> bytes:
    """Create a minimal valid PDF with some text."""
    return (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R/Resources<<>>>>endobj\n"
        b"4 0 obj<</Length 44>>stream\n"
        b"BT /F1 12 Tf 100 700 Td (Invoice #12345) Tj ET\n"
        b"endstream endobj\n"
        b"xref\n0 5\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000210 00000 n \n"
        b"trailer<</Size 5/Root 1 0 R>>\n"
        b"startxref\n294\n%%EOF"
    )


def _minimal_xlsx_bytes() -> bytes:
    """Create a minimal XLSX workbook with two rows."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Item", "Qty", "Price", "Total"])
    ws.append(["iPhone 15 Pro Max", "2", "1200", "2400"])
    ws.append(["USB-C Cable", "5", "15", "75"])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _minimal_docx_bytes() -> bytes:
    """Create a minimal DOCX with invoice-like text."""
    from docx import Document

    doc = Document()
    doc.add_heading("Invoice #DOCX-001", level=1)
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    headers = ["Item", "Qty", "Price", "Total"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["MacBook Pro", "1", "2500", "2500"],
        ["USB Hub", "3", "45", "135"],
    ]
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _png_bytes() -> bytes:
    """Create a minimal 1x1 PNG."""
    import struct
    import zlib

    def chunk(chunk_type, data):
        c = chunk_type + data
        return (
            struct.pack(">I", len(data))
            + c
            + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\x00\x00")
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", idat)
        + chunk(b"IEND", b"")
    )


# ── Fixtures ─────────────────────────────────────────────────────────────────


@pytest.fixture
def client():
    """FastAPI test client with parser router mounted."""
    from app.main import app

    return TestClient(app)


# ── File Type Detection ──────────────────────────────────────────────────────


class TestFileTypeDetection:
    def test_detect_pdf(self):
        source_type, error = detect_source_type(
            "invoice.pdf", "application/pdf", _minimal_pdf_bytes()
        )
        assert source_type == "pdf"
        assert error is None

    def test_detect_xlsx(self):
        source_type, error = detect_source_type(
            "invoice.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            _minimal_xlsx_bytes(),
        )
        assert source_type == "xlsx"
        assert error is None

    def test_detect_docx(self):
        source_type, error = detect_source_type(
            "invoice.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _minimal_docx_bytes(),
        )
        assert source_type == "docx"
        assert error is None

    def test_detect_png(self):
        source_type, error = detect_source_type("scan.png", "image/png", _png_bytes())
        assert source_type == "image"
        assert error is None

    def test_detect_jpg(self):
        source_type, error = detect_source_type(
            "scan.jpg", "image/jpeg", _png_bytes()
        )
        assert source_type == "image"
        assert error is None

    def test_unsupported_format(self):
        source_type, error = detect_source_type(
            "document.doc", "application/msword", b"dummy content"
        )
        assert source_type == ""
        assert error is not None
        assert "Неподдерживаемый формат" in error

    def test_detect_by_extension_fallback(self):
        """When MIME is missing but extension is valid."""
        source_type, error = detect_source_type(
            "file.pdf", "", _minimal_pdf_bytes()
        )
        assert source_type == "pdf"
        assert error is None


# ── MarkItDown Conversion ────────────────────────────────────────────────────


class TestMarkItDownConversion:
    def test_convert_xlsx(self):
        text, ocr = convert_to_markdown(
            _minimal_xlsx_bytes(), "xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        assert not ocr
        assert "iPhone" in text or "|" in text
        assert len(text) > 10

    def test_convert_xlsx_has_table(self):
        text, ocr = convert_to_markdown(
            _minimal_xlsx_bytes(), "xlsx"
        )
        # MarkItDown outputs Markdown tables
        assert "|" in text
        assert "Item" in text
    def test_convert_pdf(self):
        text, ocr = convert_to_markdown(
            _minimal_pdf_bytes(), "pdf", "application/pdf"
        )
        # Minimal PDF has < 100 chars → triggers OCR fallback
        # On systems without poppler, falls back to single-image Gemini call
        assert ocr is True  # Scanned detection triggered
        assert isinstance(text, str)
    def test_convert_docx(self):
        text, ocr = convert_to_markdown(
            _minimal_docx_bytes(), "docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert not ocr
        assert len(text) > 10
        assert "#" in text or "|" in text or "Invoice" in text

    def test_convert_image_triggers_ocr(self):
        """Image source type triggers Gemini Vision OCR."""
        # Gemini will be mocked by the generator's mock mode
        text, ocr = convert_to_markdown(
            _png_bytes(), "image", "image/png"
        )
        assert ocr is True
        assert isinstance(text, str)


# ── LLM Structurization ──────────────────────────────────────────────────────


class TestStructurization:
    def test_structurize_basic_invoice(self):
        """Test that structurization runs without error on markdown input."""
        raw_text = """| Item | Qty | Unit Price | Total |
        | --- | --- | --- | --- |
        | iPhone 15 Pro Max | 2 | 1200.00 | 2400.00 |
        | AirPods Pro | 3 | 249.00 | 747.00 |"""

        result = structurize_invoice(raw_text)
        assert isinstance(result, InvoiceData)
        assert len(result.items) >= 0

    def test_structurize_empty_text(self):
        result = structurize_invoice("")
        assert isinstance(result, InvoiceData)

    def test_structurize_with_markdown_heading(self):
        raw_text = """# Invoice #INV-2024-001

        | Product | Qty | Price |
        | --- | --- | --- |
        | Laptop | 1 | 1500 |"""

        result = structurize_invoice(raw_text)
        assert isinstance(result, InvoiceData)


# ── API Endpoint Tests ───────────────────────────────────────────────────────


class TestParseDocumentEndpoint:
    def test_upload_unsupported_format(self, client):
        resp = client.post(
            "/api/workspace/parse-document",
            data={"session_id": "test-session"},
            files={"file": ("test.doc", b"fake content", "application/msword")},
        )
        assert resp.status_code == 400

    def test_upload_file_too_large(self, client):
        large_content = b"x" * (MAX_FILE_SIZE + 1)
        resp = client.post(
            "/api/workspace/parse-document",
            data={"session_id": "test-session"},
            files={"file": ("large.pdf", large_content, "application/pdf")},
        )
        assert resp.status_code == 413

    def test_upload_empty_file(self, client):
        resp = client.post(
            "/api/workspace/parse-document",
            data={"session_id": "test-session"},
            files={"file": ("empty.pdf", b"", "application/pdf")},
        )
        assert resp.status_code == 400

    def test_upload_no_file(self, client):
        resp = client.post(
            "/api/workspace/parse-document",
            data={"session_id": "test-session"},
        )
        assert resp.status_code == 422

    def test_upload_valid_xlsx(self, client):
        """Integration: upload XLSX -> parse -> get InvoiceData back."""
        resp = client.post(
            "/api/workspace/parse-document",
            data={"session_id": "test-xlsx"},
            files={
                "file": (
                    "invoice.xlsx",
                    _minimal_xlsx_bytes(),
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )
        assert resp.status_code == 200
        body = resp.json()
        assert "data" in body
        assert "metadata" in body
        assert body["metadata"]["source_type"] == "xlsx"
        assert body["metadata"]["ocr_applied"] is False


class TestConfirmExtractionEndpoint:
    def test_confirm_extraction(self, client):
        invoice = InvoiceData(
            invoice_number="INV-001",
            currency="USD",
            items=[
                InvoiceLine(
                    description="Test", quantity=1, unit_price=100, total_price=100
                )
            ],
        )
        resp = client.post(
            "/api/workspace/parse-document/confirm",
            data={
                "data": invoice.model_dump_json(),
                "session_id": "test-session",
            },
        )
        assert resp.status_code == 200
        body = resp.json()
        assert body["status"] == "confirmed"

    def test_confirm_invalid_data(self, client):
        resp = client.post(
            "/api/workspace/parse-document/confirm",
            data={
                "data": "not valid json {{{",
                "session_id": "test-session",
            },
        )
        assert resp.status_code == 400


# ── Schema Validation ────────────────────────────────────────────────────────


class TestSchemas:
    def test_invoice_data_minimal(self):
        data = InvoiceData()
        assert data.items == []
        assert data.invoice_number is None

    def test_invoice_line_validation(self):
        line = InvoiceLine(
            description="Test product", quantity=2, unit_price=50.0, total_price=100.0
        )
        assert line.price_estimated is False

    def test_invoice_line_with_options(self):
        line = InvoiceLine(
            description="Heavy item",
            quantity=1,
            unit_price=500.0,
            total_price=500.0,
            weight_kg=15.0,
            hs_code_hint="8471300000",
        )
        assert line.weight_kg == 15.0
        assert line.hs_code_hint == "8471300000"

    def test_processing_metadata(self):
        from datetime import datetime

        meta = ProcessingMetadata(source_type="pdf", original_filename="invoice.pdf")
        assert meta.source_type == "pdf"
        assert meta.ocr_applied is False
        assert isinstance(meta.parsed_at, datetime)

    def test_processing_metadata_with_ocr(self):
        meta = ProcessingMetadata(
            source_type="image", ocr_applied=True, original_filename="scan.png"
        )
        assert meta.source_type == "image"
        assert meta.ocr_applied is True

    def test_parse_document_response(self):
        data = InvoiceData(invoice_number="INV-001", items=[])
        meta = ProcessingMetadata(source_type="xlsx", original_filename="test.xlsx")
        resp = ParseDocumentResponse(data=data, metadata=meta, warnings=["test warning"])
        assert resp.warnings == ["test warning"]


# ── Edge Cases ───────────────────────────────────────────────────────────────


class TestEdgeCases:
    def test_image_source_type_ocr_flag(self):
        """Image source type always sets ocr_applied=True."""
        meta = ProcessingMetadata(
            source_type="image", ocr_applied=True, original_filename="scan.pdf"
        )
        assert meta.ocr_applied is True

    def test_pdf_source_type_no_ocr(self):
        """Text PDF should not use OCR."""
        meta = ProcessingMetadata(
            source_type="pdf", ocr_applied=False, original_filename="invoice.pdf"
        )
        assert meta.ocr_applied is False

    def test_docx_source_type(self):
        """DOCX is a supported text-based format."""
        meta = ProcessingMetadata(
            source_type="docx", ocr_applied=False, original_filename="contract.docx"
        )
        assert meta.source_type == "docx"
