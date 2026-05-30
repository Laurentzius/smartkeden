import os
from app.core.documents.generator import (
    DocumentGenerator,
    CustomsInvoiceSchema,
    InvoiceItemSchema,
    SupplyAgreementSchema,
)


def test_generate_invoice_excel_with_schema(tmp_path):
    output_file = str(tmp_path / "test_invoice_schema.xlsx")
    schema = CustomsInvoiceSchema(
        seller_name="Alibaba Export Ltd.",
        buyer_name="SmartKeden Almaty LLP",
        incoterms="FCA Ningbo",
        items=[
            InvoiceItemSchema(
                name="Industrial Controller Parts",
                hs_code="8543709000",
                qty=12.0,
                unit="pcs",
                price=150.0,
            )
        ],
    )
    result_path = DocumentGenerator.generate_invoice_excel(schema, output_file)
    assert os.path.exists(result_path)
    # Structural and Content Assertions using openpyxl
    import openpyxl

    wb = openpyxl.load_workbook(result_path)
    assert "Commercial Invoice" in wb.sheetnames
    ws = wb["Commercial Invoice"]
    # Assert Title
    assert ws["A1"].value == "COMMERCIAL INVOICE / КОММЕРЧЕСКИЙ ИНВОЙС"
    # Assert Parties
    assert ws["A4"].value == "Alibaba Export Ltd."
    assert ws["A9"].value == "SmartKeden Almaty LLP"


def test_generate_invoice_excel_with_legacy_dict(tmp_path):
    output_file = str(tmp_path / "test_invoice_legacy.xlsx")
    legacy_data = {
        "seller_name": "Legacy Alibaba Ltd.",
        "buyer_name": "Legacy SmartKeden LLP",
        "incoterms": "CIF Almaty",
        "items": [
            {
                "name": "Legacy Parts",
                "hs_code": "8543709000",
                "qty": 100,
                "unit": "pcs",
                "price": 5.0,
            }
        ],
    }
    result_path = DocumentGenerator.generate_invoice_excel(legacy_data, output_file)
    assert os.path.exists(result_path)
    # Verify legacy dict contents are parsed and written properly
    import openpyxl

    wb = openpyxl.load_workbook(result_path)
    ws = wb["Commercial Invoice"]
    assert ws["A4"].value == "Legacy Alibaba Ltd."
    assert ws["A9"].value == "Legacy SmartKeden LLP"


def test_generate_contract_word_with_schema(tmp_path):
    output_file = str(tmp_path / "test_contract_schema.docx")
    schema = SupplyAgreementSchema(
        contract_no="SK-2026-99",
        contract_date="28 мая 2026 г.",
        seller_name="Alibaba Export Ltd.",
        buyer_name="SmartKeden Almaty LLP",
        incoterms="DDP Astana",
    )
    result_path = DocumentGenerator.generate_contract_word(schema, output_file)
    assert os.path.exists(result_path)
    # Structural and Content Assertions using python-docx
    import docx

    doc = docx.Document(result_path)
    # Verify Heading
    assert doc.paragraphs[0].text == "ДОГОВОР ПОСТАВКИ № SK-2026-99"
    # Verify Date and Incoterms in content
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "28 мая 2026 г." in full_text
    assert "DDP Astana" in full_text


def test_generate_contract_word_with_legacy_dict(tmp_path):
    output_file = str(tmp_path / "test_contract_legacy.docx")
    legacy_data = {
        "contract_no": "SK-LEGACY-01",
        "contract_date": "15 января 2024 г.",
        "seller_name": "Legacy Seller Corp",
        "buyer_name": "Legacy Buyer LLP",
        "incoterms": "CIP Almaty",
    }
    result_path = DocumentGenerator.generate_contract_word(legacy_data, output_file)
    assert os.path.exists(result_path)
    import docx

    doc = docx.Document(result_path)
    assert doc.paragraphs[0].text == "ДОГОВОР ПОСТАВКИ № SK-LEGACY-01"
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "15 января 2024 г." in full_text
    assert "CIP Almaty" in full_text
