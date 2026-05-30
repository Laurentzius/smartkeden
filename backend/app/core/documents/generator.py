import logging
import os
from typing import Dict, Any, Optional, List, Union
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)


class InvoiceItemSchema(BaseModel):
    name: str = Field(..., description="Description of the goods")
    hs_code: str = Field(..., description="10-digit HS code (ТН ВЭД)")
    qty: float = Field(..., description="Quantity of the goods")
    unit: str = Field("pcs", description="Unit of measurement (pcs, kg, meters, etc.)")
    price: float = Field(..., description="Unit price in transaction currency")


class CustomsInvoiceSchema(BaseModel):
    seller_name: str = Field(..., description="Name of the selling company (Vendor)")
    buyer_name: str = Field(..., description="Name of the buying company (Importer)")
    incoterms: str = Field("FCA Shenzhen", description="Incoterms delivery condition")
    items: List[InvoiceItemSchema] = Field(..., description="List of invoiced goods")


class SupplyAgreementSchema(BaseModel):
    contract_no: str = Field(..., description="Unique agreement contract number")
    contract_date: str = Field(..., description="Date of contract signing")
    seller_name: str = Field(..., description="Selling company name")
    buyer_name: str = Field(..., description="Buying company name")
    incoterms: str = Field("DDP Almaty", description="Incoterms delivery condition")


class DocumentGenerator:
    """
    Generates formal trade documents:
    - Invoices & Specifications (Excel / .xlsx)
    - Supply Agreements / Договоры поставки (Word / .docx)
    - Analytical party tables and Customs Declarations summaries (PDF)
    """

    @staticmethod
    def generate_invoice_excel(
        data: Union[CustomsInvoiceSchema, Dict[str, Any]], output_path: str
    ) -> str:
        """
        Generates a standard commercial invoice in Excel (.xlsx) format using CustomsInvoiceSchema.
        """
        if isinstance(data, dict):
            # Support legacy dict structures
            resolved_items = []
            for item in data.get("items", []):
                resolved_items.append(
                    InvoiceItemSchema(
                        name=item.get("name", "Sample Goods"),
                        hs_code=item.get("hs_code", "8543709000"),
                        qty=float(item.get("qty", 100)),
                        unit=item.get("unit", "pcs"),
                        price=float(item.get("price", 10.0)),
                    )
                )
            schema = CustomsInvoiceSchema(
                seller_name=data.get("seller_name", "Vendor Inc."),
                buyer_name=data.get("buyer_name", "Kazakhstan Importer LLP"),
                incoterms=data.get("incoterms", "FCA Shenzhen"),
                items=resolved_items,
            )
        else:
            schema = data

        logger.info(f"Generating commercial invoice Excel at {output_path}")
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Commercial Invoice"
            # Ensure grid lines are visible
            ws.views.sheetView[0].showGridLines = True
            # --- Styles ---
            font_title = Font(name="Calibri", size=16, bold=True, color="1F4E78")
            font_section = Font(name="Calibri", size=10, bold=True, color="1F4E78")
            font_header = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
            font_data = Font(name="Calibri", size=10, bold=False)
            font_total = Font(name="Calibri", size=10, bold=True)
            font_small_italic = Font(
                name="Calibri", size=8, italic=True, color="595959"
            )
            fill_header = PatternFill(
                start_color="1F4E78", end_color="1F4E78", fill_type="solid"
            )
            fill_zebra = PatternFill(
                start_color="F2F5F8", end_color="F2F5F8", fill_type="solid"
            )
            fill_total = PatternFill(
                start_color="E9EEF4", end_color="E9EEF4", fill_type="solid"
            )
            align_center = Alignment(
                horizontal="center", vertical="center", wrap_text=True
            )
            align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
            align_right = Alignment(horizontal="right", vertical="center")
            border_thin = Side(border_style="thin", color="D9D9D9")
            border_medium = Side(border_style="medium", color="1F4E78")
            border_double = Side(border_style="double", color="1F4E78")
            box_border = Border(
                left=border_thin, right=border_thin, top=border_thin, bottom=border_thin
            )
            header_border = Border(
                left=border_thin,
                right=border_thin,
                top=border_medium,
                bottom=border_medium,
            )
            total_border = Border(top=border_thin, bottom=border_double)
            # --- Page Header (Title & Metadata) ---
            ws.merge_cells("A1:G1")
            ws["A1"] = "COMMERCIAL INVOICE / КОММЕРЧЕСКИЙ ИНВОЙС"
            ws["A1"].font = font_title
            ws["A1"].alignment = align_center
            ws.row_dimensions[1].height = 35
            # Metadata Block
            ws["E3"] = "Invoice No / Инвойс №:"
            ws["E3"].font = font_total
            ws["F3"] = "INV-2026-089"
            ws["F3"].font = font_data
            ws["E4"] = "Date / Дата:"
            ws["E4"].font = font_total
            ws["F4"] = "28.05.2026"
            ws["F4"].font = font_data
            ws["E5"] = "Contract / Контракт:"
            ws["E5"].font = font_total
            ws["F5"] = "KED-2026/089"
            ws["F5"].font = font_data
            # --- Parties Block (Seller vs Buyer) ---
            # Seller block
            ws.merge_cells("A3:C3")
            ws["A3"] = "SELLER / ПРОДАВЕЦ"
            ws["A3"].font = font_section
            ws["A3"].fill = PatternFill(
                start_color="E9EEF4", end_color="E9EEF4", fill_type="solid"
            )
            ws.merge_cells("A4:C4")
            ws["A4"] = schema.seller_name
            ws["A4"].font = font_total
            ws["A4"].alignment = align_left
            ws.merge_cells("A5:C6")
            ws["A5"] = (
                "Industrial District, Nanshan, Shenzhen, Guangdong, China\nTax ID: 91440300MA5EXXXX"
            )
            ws["A5"].font = font_data
            ws["A5"].alignment = align_left
            # Buyer block
            ws.merge_cells("A8:C8")
            ws["A8"] = "BUYER / ПОКУПАТЕЛЬ"
            ws["A8"].font = font_section
            ws["A8"].fill = PatternFill(
                start_color="E9EEF4", end_color="E9EEF4", fill_type="solid"
            )
            ws.merge_cells("A9:C9")
            ws["A9"] = schema.buyer_name
            ws["A9"].font = font_total
            ws["A9"].alignment = align_left
            ws.merge_cells("A10:C11")
            ws["A10"] = (
                "Республика Казахстан, г. Алматы, Медеуский р-н, пр. Достык, 120\nБИН: 120440056123, Тел: +7 (727) 333-22-11"
            )
            ws["A10"].font = font_data
            ws["A10"].alignment = align_left
            # Delivery & Origin info on the right
            ws["E8"] = "Incoterms / Инкотермс:"
            ws["E8"].font = font_total
            ws["F8"] = schema.incoterms
            ws["F8"].font = font_data
            ws["E9"] = "Origin / Страна происх.:"
            ws["E9"].font = font_total
            ws["F9"] = "CHINA (CN) / КИТАЙ"
            ws["F9"].font = font_data
            ws["E10"] = "Currency / Валюта:"
            ws["E10"].font = font_total
            ws["F10"] = (
                data.get("items", [{}])[0].get("currency", "USD")
                if isinstance(data, dict)
                else "USD"
            )
            ws["F10"].font = font_data
            # --- Table Headers (Row 13) ---
            headers = [
                "Item\n№",
                "Description of Goods / Описание товара",
                "HS Code\nКод ТН ВЭД",
                "Qty\nКол-во",
                "Unit\nЕд.",
                "Unit Price\nЦена за ед.",
                "Total Amount\nСумма",
            ]
            ws.row_dimensions[13].height = 28
            for col_idx, h in enumerate(headers, 1):
                cell = ws.cell(row=13, column=col_idx, value=h)
                cell.font = font_header
                cell.fill = fill_header
                cell.alignment = align_center
                cell.border = header_border
            # --- Table Data ---
            row_idx = 14
            total_sum = 0.0
            for idx, item in enumerate(schema.items, 1):
                ws.row_dimensions[row_idx].height = 22
                c1 = ws.cell(row=row_idx, column=1, value=idx)
                c1.font = font_data
                c1.alignment = align_center
                c1.border = box_border
                c2 = ws.cell(row=row_idx, column=2, value=item.name)
                c2.font = font_data
                c2.alignment = align_left
                c2.border = box_border
                c3 = ws.cell(row=row_idx, column=3, value=item.hs_code)
                c3.font = font_data
                c3.alignment = align_center
                c3.border = box_border
                c4 = ws.cell(row=row_idx, column=4, value=item.qty)
                c4.font = font_data
                c4.alignment = align_center
                c4.border = box_border
                c4.number_format = "#,##0.00"
                c5 = ws.cell(row=row_idx, column=5, value=item.unit)
                c5.font = font_data
                c5.alignment = align_center
                c5.border = box_border
                c6 = ws.cell(row=row_idx, column=6, value=item.price)
                c6.font = font_data
                c6.alignment = align_right
                c6.border = box_border
                c6.number_format = "#,##0.00"
                amt = item.qty * item.price
                total_sum += amt
                c7 = ws.cell(row=row_idx, column=7, value=amt)
                c7.font = font_data
                c7.alignment = align_right
                c7.border = box_border
                c7.number_format = "#,##0.00"
                # Zebra striping
                if idx % 2 == 0:
                    for col_c in range(1, 8):
                        ws.cell(row=row_idx, column=col_c).fill = fill_zebra
                row_idx += 1
            # --- Total Row ---
            ws.row_dimensions[row_idx].height = 24
            ws.merge_cells(
                start_row=row_idx, start_column=1, end_row=row_idx, end_column=6
            )
            total_label_cell = ws.cell(
                row=row_idx, column=1, value="TOTAL AMOUNT / ИТОГО К ОПЛАТЕ:"
            )
            total_label_cell.font = font_total
            total_label_cell.alignment = Alignment(
                horizontal="right", vertical="center"
            )
            # Apply borders & fills to merged total label cells
            for col_c in range(1, 7):
                ws.cell(row=row_idx, column=col_c).border = Border(
                    top=border_thin, bottom=border_thin
                )
                ws.cell(row=row_idx, column=col_c).fill = fill_total
            total_val_cell = ws.cell(row=row_idx, column=7, value=total_sum)
            total_val_cell.font = font_total
            total_val_cell.fill = fill_total
            total_val_cell.alignment = align_right
            total_val_cell.border = total_border
            total_val_cell.number_format = "#,##0.00"
            # --- Signatures Block ---
            row_idx += 3
            ws.merge_cells(
                start_row=row_idx, start_column=1, end_row=row_idx, end_column=3
            )
            ws.cell(
                row=row_idx, column=1, value="For Seller / От имени Продавца:"
            ).font = font_total
            ws.merge_cells(
                start_row=row_idx, start_column=5, end_row=row_idx, end_column=7
            )
            ws.cell(
                row=row_idx, column=5, value="For Buyer / От имени Покупателя:"
            ).font = font_total
            row_idx += 2
            ws.cell(
                row=row_idx, column=1, value="___________________________________"
            ).font = font_data
            ws.cell(
                row=row_idx, column=5, value="___________________________________"
            ).font = font_data
            row_idx += 1
            ws.cell(
                row=row_idx, column=1, value="Signature (Подпись) / L.S. (М.П.)"
            ).font = font_small_italic
            ws.cell(
                row=row_idx, column=5, value="Signature (Подпись) / L.S. (М.П.)"
            ).font = font_small_italic
            # --- Column Width Adjustments ---
            column_widths = {
                "A": 6,  # No
                "B": 42,  # Description
                "C": 15,  # HS Code
                "D": 10,  # Qty
                "E": 8,  # Unit
                "F": 15,  # Price
                "G": 18,  # Total Amount
            }
            for col, width in column_widths.items():
                ws.column_dimensions[col].width = width
            # Save file
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            wb.save(output_path)
            return output_path
        except ImportError:
            logger.warning("openpyxl not installed, writing basic dummy text instead")
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(f"Invoice Schema: {schema.model_dump_json()}")
            return output_path

    @staticmethod
    def generate_contract_word(
        data: Union[SupplyAgreementSchema, Dict[str, Any]], output_path: str
    ) -> str:
        """
        Generates a standard foreign trade supply agreement in Word (.docx) format using SupplyAgreementSchema.
        """
        if isinstance(data, dict):
            schema = SupplyAgreementSchema(
                contract_no=data.get("contract_no", "2026/01"),
                contract_date=data.get("contract_date", "27 мая 2026 г."),
                seller_name=data.get("seller_name", "Vendor Inc."),
                buyer_name=data.get("buyer_name", "Kazakhstan Importer LLP"),
                incoterms=data.get("incoterms", "DDP Almaty"),
            )
        else:
            schema = data

        logger.info(f"Generating supply contract Word document at {output_path}")
        try:
            import docx
            from docx import Document

            doc = docx.Document()
            doc.add_heading("ДОГОВОР ПОСТАВКИ № " + schema.contract_no, level=1)

            p = doc.add_paragraph()
            p.add_run("г. Алматы\t\t\t\t\tДата: ").bold = True
            p.add_run(schema.contract_date)

            doc.add_heading("1. ПРЕДМЕТ ДОГОВОРА", level=2)
            doc.add_paragraph(
                f"Продавец обязуется поставить, а Покупатель принять и оплатить товар в соответствии "
                f"со спецификациями к настоящему договору на условиях {schema.incoterms}."
            )

            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            return output_path
        except ImportError:
            logger.warning("docx not installed, writing basic dummy text instead")
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(f"Contract Schema: {schema.model_dump_json()}")
            return output_path
